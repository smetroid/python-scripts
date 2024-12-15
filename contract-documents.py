from docx import Document
import re
from datetime import datetime
import os
from os.path import expanduser

def extract_date_from_filename(filename):
    """
    Extract a date from the filename using a regex pattern.
    Adjust the regex to match the date format in your filenames.
    """
    # Example regex for dates in format MM-DD-YY
    date_pattern = r"(\d{2}-\d{2}-\d{2})"
    match = re.search(date_pattern, filename)
    if match:
        return datetime.strptime(match.group(1), "%m-%d-%y")
    return None

def combine_documents_in_order(folder_path, base_file, output_file):
    """
    Combine Word documents in chronological order based on dates in filenames,
    appending them to a specified base document.

    :param folder_path: Path to the folder containing the Word documents.
    :param base_file: Path to the base Word document.
    :param output_file: Name of the combined output file.
    """
    # Expand user path if it contains ~
    folder_path = expanduser(folder_path)
    base_file = expanduser(base_file)
    output_file = expanduser(output_file)

    # List all .docx files in the folder
    files = [f for f in os.listdir(folder_path) if f.endswith(".docx")]

    # Extract dates and sort files by date
    files_with_dates = []
    for file in files:
        file_path = os.path.join(folder_path, file)
        file_date = extract_date_from_filename(file)
        if file_date:
            files_with_dates.append((file_date, file_path))

    # Sort by date
    files_with_dates.sort(key=lambda x: x[0])

    # Combine the files
    if not files_with_dates:
        print("No valid files with dates found.")
        return

    combined_doc = Document(base_file)  # Start with the specified base document

    for _, file_path in files_with_dates:
        # Add the filename as a header
        file_name = os.path.basename(file_path)
        combined_doc.add_paragraph(f"--- {file_name} ---", style="Heading 2")

        # Append the content of the file after the header
        doc = Document(file_path)
        for para in doc.paragraphs:
            combined_doc.add_paragraph(para.text)

    # Save the combined document
    combined_doc.save(output_file)
    print(f"Documents combined and saved as {output_file}")

# Example usage
folder = "~/Downloads/contracts/"  # Folder containing Word documents
base = "~/Downloads/contracts/Contracts Outline 12-13-24.docx"  # Base document to append to
output = "./combined.docx"  # Output file name
combine_documents_in_order(folder, base, output)
